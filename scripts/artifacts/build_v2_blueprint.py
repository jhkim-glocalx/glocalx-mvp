from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/jaehun/01_projects/glocalx-mvp")
SOURCE = ROOT / "docs/v2/glocalx-v2-blueprint.md"
OUTPUT = ROOT / "outputs/v2/glocalx-v2-blueprint.docx"

COLORS = {
    "canvas": "0C0B10",
    "surface": "FBF9F6",
    "card": "FFFFFF",
    "ink": "191720",
    "ink_soft": "48424F",
    "muted": "938C9C",
    "line": "ECE7EF",
    "accent": "FF6A3D",
    "accent_soft": "FFF1EC",
    "mint": "15BD97",
    "mint_soft": "E6F8F2",
    "blue": "3D6BFF",
}


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, attrs in edges.items():
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        for key, value in attrs.items():
            edge.set(qn(f"w:{key}"), str(value))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_width(table, width_inches: float) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tbl_w.set(qn("w:type"), "dxa")


def set_paragraph_border(paragraph, color: str, size: str = "8", space: str = "1") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), COLORS["blue"])
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.extend([color, underline])
    run.append(run_props)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_PATTERN = re.compile(r"(\*\*.+?\*\*|\[[^\]]+\]\([^)]+\))")


def add_inline(paragraph, text: str, color: str | None = None) -> None:
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            if color:
                run.font.color.rgb = rgb(color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            if color:
                run.font.color.rgb = rgb(color)
        elif token.startswith("["):
            link_match = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link_match:
                add_hyperlink(paragraph, link_match.group(1), link_match.group(2))
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        if color:
            run.font.color.rgb = rgb(color)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(COLORS["ink"])
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(10.5)
        style.font.color.rgb = rgb(COLORS["ink"])
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.left_indent = Inches(0.23)
        style.paragraph_format.first_line_indent = Inches(-0.15)

    heading_specs = {
        "Title": (34, COLORS["card"], 0, 10),
        "Subtitle": (17, COLORS["card"], 0, 18),
        "Heading 1": (17, COLORS["ink"], 18, 7),
        "Heading 2": (13, COLORS["accent"], 12, 5),
        "Heading 3": (11.5, COLORS["ink_soft"], 9, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    styles["Heading 1"].paragraph_format.page_break_before = False


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.66)
    section.left_margin = Inches(0.76)
    section.right_margin = Inches(0.76)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.text = "GLOCALX  /  V2 PRODUCT & OPERATIONS BLUEPRINT"
    p.style = doc.styles["Normal"]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(7.5)
        run.font.bold = True
        run.font.color.rgb = rgb(COLORS["muted"])
    set_paragraph_border(p, COLORS["line"], size="5")

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(2)
    prefix = fp.add_run("CONFIDENTIAL PLANNING BASELINE   •   16 JULY 2026   •   ")
    prefix.font.size = Pt(7)
    prefix.font.color.rgb = rgb(COLORS["muted"])
    add_page_number(fp)

    core = doc.core_properties
    core.title = "GlocalX V2 Product, Operations, and Delivery Blueprint"
    core.subject = "Customer portal, admin operations dashboard, managed marketing, and delivery plan"
    core.author = "GlocalX"
    core.language = "en-US"
    core.keywords = "GlocalX, V2, product, operations, architecture, Google Business Profile, customer support"
    core.comments = "Generated from the reviewed Markdown planning baseline."


def add_cover(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 6.9)
    cell = table.cell(0, 0)
    set_cell_shading(cell, COLORS["canvas"])
    set_cell_margins(cell, top=360, start=380, bottom=320, end=380)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    row = table.rows[0]
    row.height = Inches(9.15)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    eyebrow = cell.paragraphs[0]
    eyebrow.paragraph_format.space_after = Pt(26)
    run = eyebrow.add_run("GLOCALX  /  VERSION 2")
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = rgb(COLORS["accent"])

    title = cell.add_paragraph()
    title.style = doc.styles["Title"]
    title.paragraph_format.space_after = Pt(12)
    title.add_run("The operating system\nfor managed local growth")

    subtitle = cell.add_paragraph()
    subtitle.style = doc.styles["Subtitle"]
    subtitle.add_run("Product, operations, architecture, and delivery blueprint")

    rule = cell.add_paragraph()
    rule.paragraph_format.space_before = Pt(10)
    rule.paragraph_format.space_after = Pt(22)
    set_paragraph_border(rule, COLORS["accent"], size="26", space="4")

    thesis = cell.add_paragraph()
    thesis.paragraph_format.space_after = Pt(22)
    thesis_run = thesis.add_run(
        "Connect the store. Send images and intent. "
        "GlocalX produces, the customer approves, and the platform publishes."
    )
    thesis_run.font.name = "Arial"
    thesis_run.font.size = Pt(14)
    thesis_run.font.bold = True
    thesis_run.font.color.rgb = rgb(COLORS["card"])

    meta = cell.add_paragraph()
    meta.paragraph_format.space_before = Pt(20)
    add_inline(
        meta,
        "Prepared 16 July 2026\n"
        "Audience: founders, investors, product, design, engineering, customer success, and operations\n"
        "Status: implementation-ready planning baseline",
        COLORS["muted"],
    )
    for run in meta.runs:
        run.font.size = Pt(9)

    doc.add_page_break()


def add_contents(doc: Document, headings: list[str]) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("Contents")
    intro = doc.add_paragraph(
        "This document is organized from strategic decision through implementation, release, and investor evidence."
    )
    intro.paragraph_format.space_after = Pt(10)

    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 6.85)
    for index, heading in enumerate(headings, 1):
        row = table.add_row()
        prevent_row_split(row)
        num_cell, title_cell = row.cells
        num_cell.width = Inches(0.52)
        title_cell.width = Inches(6.33)
        set_cell_margins(num_cell, top=55, start=60, bottom=55, end=50)
        set_cell_margins(title_cell, top=55, start=90, bottom=55, end=60)
        if index % 2 == 0:
            set_cell_shading(num_cell, COLORS["surface"])
            set_cell_shading(title_cell, COLORS["surface"])
        number_p = num_cell.paragraphs[0]
        number_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        nr = number_p.add_run(f"{index:02d}")
        nr.font.bold = True
        nr.font.size = Pt(8.5)
        nr.font.color.rgb = rgb(COLORS["accent"])
        title_p = title_cell.paragraphs[0]
        tr = title_p.add_run(heading)
        tr.font.size = Pt(9)
        tr.font.color.rgb = rgb(COLORS["ink"])
def add_callout(doc: Document, title: str, body: str, fill: str, accent: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 6.75)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=130, start=170, bottom=130, end=170)
    set_cell_border(
        cell,
        start={"val": "single", "sz": "18", "color": accent},
        top={"val": "nil"},
        bottom={"val": "nil"},
        end={"val": "nil"},
    )
    title_p = cell.paragraphs[0]
    title_run = title_p.add_run(title.upper())
    title_run.font.size = Pt(8.5)
    title_run.font.bold = True
    title_run.font.color.rgb = rgb(accent)
    body_p = cell.add_paragraph()
    body_p.paragraph_format.space_after = Pt(0)
    add_inline(body_p, body)


def add_service_loop(doc: Document) -> None:
    label = doc.add_paragraph()
    label.paragraph_format.space_before = Pt(8)
    lr = label.add_run("SERVICE LOOP AT A GLANCE")
    lr.font.size = Pt(8)
    lr.font.bold = True
    lr.font.color.rgb = rgb(COLORS["muted"])

    steps = [
        ("01", "Connect", "Google identity + GBP"),
        ("02", "Submit", "Images + short intent"),
        ("03", "Produce", "GlocalX operations"),
        ("04", "Approve", "Customer go / no-go"),
        ("05", "Publish", "Channel jobs + recovery"),
    ]
    table = doc.add_table(rows=1, cols=len(steps))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 6.8)
    for idx, (number, title, subtitle) in enumerate(steps):
        cell = table.cell(0, idx)
        set_cell_shading(cell, COLORS["surface"] if idx % 2 == 0 else COLORS["accent_soft"])
        set_cell_margins(cell, top=120, start=85, bottom=120, end=85)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        n = p.add_run(number + "\n")
        n.font.size = Pt(8)
        n.font.bold = True
        n.font.color.rgb = rgb(COLORS["accent"])
        t = p.add_run(title + "\n")
        t.font.size = Pt(9.5)
        t.font.bold = True
        t.font.color.rgb = rgb(COLORS["ink"])
        s = p.add_run(subtitle)
        s.font.size = Pt(7)
        s.font.color.rgb = rgb(COLORS["muted"])


def add_support_modes(doc: Document) -> None:
    modes = [
        ("AI_AUTO", "AI replies", COLORS["mint_soft"]),
        ("AI_SUGGEST", "Human reviews", COLORS["accent_soft"]),
        ("HUMAN", "Human only", "EEF2FF"),
        ("CLOSED", "Resolved", COLORS["surface"]),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 6.8)
    for idx, (mode, description, fill) in enumerate(modes):
        cell = table.cell(0, idx)
        set_cell_shading(cell, fill)
        set_cell_margins(cell, top=100, start=85, bottom=100, end=85)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mr = p.add_run(mode + "\n")
        mr.font.size = Pt(8.5)
        mr.font.bold = True
        mr.font.color.rgb = rgb(COLORS["ink"])
        dr = p.add_run(description)
        dr.font.size = Pt(7.5)
        dr.font.color.rgb = rgb(COLORS["ink_soft"])


def add_architecture(doc: Document) -> None:
    label = doc.add_paragraph()
    lr = label.add_run("TWO SURFACES, ONE OPERATING PLATFORM")
    lr.font.size = Pt(8)
    lr.font.bold = True
    lr.font.color.rgb = rgb(COLORS["muted"])

    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 6.8)
    content = [
        ("CUSTOMER PORTAL", "ADMIN DASHBOARD", COLORS["surface"], COLORS["surface"]),
        (
            "Identity • GBP connection • Intake • Approvals • Contextual chat",
            "Access • Support • Creative • Publishing • Recovery",
            COLORS["accent_soft"],
            COLORS["mint_soft"],
        ),
        (
            "SHARED DOMAIN • API • POSTGRES • OBJECT STORAGE • AUDIT",
            "WORKER • QUEUE • AI • GOOGLE GBP • INSTAGRAM • NOTIFICATIONS",
            COLORS["canvas"],
            COLORS["canvas"],
        ),
    ]
    for row_index, row_content in enumerate(content):
        row = table.rows[row_index]
        prevent_row_split(row)
        for col_index in range(2):
            text, fill = row_content[col_index], row_content[col_index + 2]
            cell = row.cells[col_index]
            set_cell_shading(cell, fill)
            set_cell_margins(cell, top=105, start=100, bottom=105, end=100)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.size = Pt(8 if row_index else 9)
            run.font.bold = True if row_index != 1 else False
            run.font.color.rgb = rgb(COLORS["card"] if row_index == 2 else COLORS["ink"])
    table.cell(2, 0).merge(table.cell(2, 1))


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    normalized = [row + [""] * (column_count - len(row)) for row in rows]
    table = doc.add_table(rows=len(normalized), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 6.82)
    font_size = 8.3 if column_count <= 4 else 7.4
    for row_index, values in enumerate(normalized):
        row = table.rows[row_index]
        prevent_row_split(row)
        if row_index == 0:
            set_repeat_table_header(row)
        for col_index, value in enumerate(values):
            cell = row.cells[col_index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_margins(cell, top=75, start=85, bottom=75, end=85)
            if row_index == 0:
                set_cell_shading(cell, COLORS["canvas"])
            elif row_index % 2 == 0:
                set_cell_shading(cell, COLORS["surface"])
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline(p, value, COLORS["card"] if row_index == 0 else COLORS["ink"])
            for run in p.runs:
                run.font.size = Pt(font_size)
                if row_index == 0:
                    run.font.bold = True
            set_cell_border(
                cell,
                bottom={"val": "single", "sz": "4", "color": COLORS["line"]},
                top={"val": "single", "sz": "4", "color": COLORS["line"]},
                start={"val": "single", "sz": "4", "color": COLORS["line"]},
                end={"val": "single", "sz": "4", "color": COLORS["line"]},
            )
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def is_separator_row(line: str) -> bool:
    stripped = line.strip().strip("|")
    cells = [cell.strip() for cell in stripped.split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        line = lines[index].strip()
        if not is_separator_row(line):
            cells = [cell.strip() for cell in line.strip("|").split("|")]
            rows.append(cells)
        index += 1
    return rows, index


def add_body(doc: Document, source: str) -> None:
    lines = source.splitlines()
    start = next(i for i, line in enumerate(lines) if line.strip() == "## Executive summary")
    page_break_sections = {
        "1. Strategic product decision",
        "4. V2 product scope",
        "10. Recommended system architecture",
        "15. Delivery roadmap",
        "20. Investor narrative",
        "Appendix A. Traceability from founder request",
    }
    index = start
    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()

        if not stripped:
            index += 1
            continue

        if stripped == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            set_paragraph_border(p, COLORS["line"], size="7")
            index += 1
            continue

        if stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            add_markdown_table(doc, rows)
            continue

        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            title = heading.group(2)
            if title in page_break_sections and len(doc.paragraphs) > 1:
                doc.add_page_break()
            style = {2: "Heading 1", 3: "Heading 2", 4: "Heading 3"}[level]
            p = doc.add_paragraph()
            p.style = doc.styles[style]
            add_inline(p, title)
            if level == 2:
                set_paragraph_border(p, COLORS["accent"], size="10", space="3")
            if title == "6. Support and AI-to-human handoff":
                add_support_modes(doc)
            elif title == "7. Managed marketing workflow":
                add_service_loop(doc)
            elif title == "10. Recommended system architecture":
                add_architecture(doc)
            index += 1
            continue

        if stripped.startswith("> "):
            add_callout(doc, "Operating thesis", stripped[2:], COLORS["accent_soft"], COLORS["accent"])
            index += 1
            continue

        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, bullet.group(1))
            index += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, numbered.group(1))
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if (
                not candidate
                or candidate == "---"
                or candidate.startswith("#")
                or candidate.startswith("|")
                or candidate.startswith("> ")
                or re.match(r"^-\s+", candidate)
                or re.match(r"^\d+\.\s+", candidate)
            ):
                break
            paragraph_lines.append(candidate)
            index += 1

        text = " ".join(paragraph_lines)
        if text.startswith("**Status:**") or text.startswith("**Prepared:**"):
            continue
        p = doc.add_paragraph()
        add_inline(p, text)


def validate_document(doc: Document) -> None:
    if len(doc.paragraphs) < 300:
        raise RuntimeError(f"Unexpectedly short document: {len(doc.paragraphs)} paragraphs")
    if len(doc.tables) < 20:
        raise RuntimeError(f"Expected extensive tables: found {len(doc.tables)}")
    heading_count = sum(1 for p in doc.paragraphs if p.style.name.startswith("Heading"))
    if heading_count < 100:
        raise RuntimeError(f"Expected complete heading hierarchy: found {heading_count}")


def main() -> int:
    if not SOURCE.exists():
        print(f"Missing source: {SOURCE}", file=sys.stderr)
        return 2

    source = SOURCE.read_text(encoding="utf-8")
    headings = [
        line.removeprefix("## ").strip()
        for line in source.splitlines()
        if line.startswith("## ") and not line.startswith("### ")
    ]

    doc = Document()
    configure_styles(doc)
    configure_document(doc)
    add_cover(doc)
    add_contents(doc, headings)
    add_body(doc, source)
    validate_document(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(
        f"created={OUTPUT}\n"
        f"bytes={OUTPUT.stat().st_size}\n"
        f"paragraphs={len(doc.paragraphs)}\n"
        f"tables={len(doc.tables)}\n"
        f"sections={len(doc.sections)}"
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
